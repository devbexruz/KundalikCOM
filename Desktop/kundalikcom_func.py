import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
import tempfile
import time
import os
import render_capcha_image
from urllib.parse import unquote
import traceback

def get_all_group(browser, school_id, set_users = False):
    text = browser.get(f"https://schools.emaktab.uz/v2/school?school={school_id}&view=classes")
    soup = BeautifulSoup(text.content, "html.parser")
    all_grades = soup.find("ul",class_="classes")
    maktab_data = dict()
    try:
        for grade in all_grades.find_all(recursive=False):
            grade_name = grade.find("strong").get_text()
            grade_data = dict() if set_users else []
            for class_ in grade.find_all("li"):
                if set_users:
                    class_id = class_.find("a").get("href").split("class=")[-1]
                    grade_data[class_id] = class_.get_text()
                else:
                    grade_data.append(class_.get_text())
            maktab_data[grade_name] = grade_data
    except:
        pass
    return maktab_data


def call_foiz(a,b):
    return (b*100)//a, ((b*1000)//a)/1000


def get_users(browser, school_id, ui, database, login, parol):
    ui.pushButton.setEnabled(False)
    database.is_get_users_thread = True
    natija = dict()
    res = browser.get(f"https://schools.emaktab.uz/v2/school?school={school_id}&view=members&group=students")
    soup = BeautifulSoup(res.content, "html.parser")
    number_pupils = int(soup.find(class_="found").find("strong").get_text().strip())



    res = browser.get(f"https://schools.emaktab.uz/v2/school?school={school_id}&view=members&group=parents")
    soup = BeautifulSoup(res.content, "html.parser")
    number_parents = int(soup.find(class_="found").find("strong").get_text().strip())

    end = -1
    all_classes = get_all_group(browser, school_id, set_users=True)
    # O'quvchilarni topib chiqish
    for num in range(1, number_pupils//10+2):
        res = browser.get(f"https://schools.emaktab.uz/v2/school?school={school_id}&view=members&group=students&page={num}")
        soup = BeautifulSoup(res.content, "html.parser")
        teen_pupils = soup.find_all(class_="tdName")
        for pupil in teen_pupils:
            sinf = "Aniqlanmagan"
            try:
                a = pupil.find("a")
                # o'quvchi ismini o'qish
                pupil_name = a.get_text().strip()
                # o'quvchi is sini o'qish
                pupil_id = a.get("href").split("user=")[-1].strip()
                natija[pupil_id] = {
                    "id": pupil_id,
                    "full_name": pupil_name,
                    "sinf": sinf,
                    "browser": requests.session(),
                    "login": None,
                    "parol": None
                }
                value, scroll = call_foiz(number_pupils+number_parents+2, len(natija))
                if database.is_get_users_thread == False:
                    # print("Finish")
                    return 0
                if value>end:
                    ui.label_3.setText(f"{value}%")
                    end = value
                    ui.label_3.setStyleSheet(f"""
                        background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:1, y2:0, stop:0 rgba(219, 255, 0, 255), stop:{scroll} rgba(0, 255, 50, 255), stop:{scroll}1 rgba(196, 196, 196, 30));
                        color: rgb(0,20, 55);
                        font: 30px "Showcard Gothic";""")
            except:
                pass

    # Ota onalarni topib chiqish
    for num in range(1, number_parents//10+2):
        while True:
            try:
                res = browser.get(f"https://schools.emaktab.uz/v2/school?school={school_id}&view=members&group=parents&page={num}")
                break
            except requests.exceptions.ConnectionError:
                continue
            except:
                break

        soup = BeautifulSoup(res.content, "html.parser")
        teen_parents = soup.find_all(class_="tdName")
        for parent in teen_parents:
            sinf = "Ota onalar"
            try:
                a = parent.find("a")

                # o'quvchi ismini o'qish
                parent_name = a.get_text().strip()

                # o'quvchi is sini o'qish
                parent_id = a.get("href").split("user=")[-1].strip()
                natija[parent_id] = {
                    "id": parent_id,
                    "full_name": parent_name,
                    "sinf": sinf,
                    "browser": requests.session(),
                    "login": None,
                    "parol": None
                }
                value, scroll = call_foiz(number_pupils+number_parents+2, len(natija))
                if database.is_get_users_thread == False:
                    # print("Finish")
                    return 0
                if value>end:
                    ui.label_3.setText(f"{value}%")
                    end = value
                    ui.label_3.setStyleSheet(f"""
                        background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:1, y2:0, stop:0 rgba(219, 255, 0, 255), stop:{scroll} rgba(0, 255, 50, 255), stop:{scroll}1 rgba(196, 196, 196, 30));
                        color: rgb(0,20, 55);
                        font: 30px "Showcard Gothic";""")
            except:
                pass
    for i in database.dict_data["all_data_logins"].keys():
        if i in natija:
            natija[i]["login"] = database.dict_data["all_data_logins"][i]["login"]
            natija[i]["parol"] = database.dict_data["all_data_logins"][i]["parol"]
    
    ui.label_3.setText("Sinflari aniqlanmoqda ...")
    for grade in all_classes.keys():
        for class_id in all_classes[grade].keys():
            r = browser.get(f"https://schools.emaktab.uz/v2/class?class={class_id}&view=members&group=students")
            class_soup = BeautifulSoup(r.content, "html.parser")
            all_users_in_class = class_soup.find_all(class_="tdAvatar")
            for class_user in all_users_in_class:
                try:
                    anker = class_user.find("a")
                    class_user_id = anker.get("href").split("user=")[-1]
                    sex = anker.get("class")[-1]
                    if class_user_id in natija:
                        natija[class_user_id]["sinf"] = all_classes[grade][class_id]
                        natija[class_user_id]["sex"] = sex
                except:
                    pass
            r = browser.get(f"https://schools.emaktab.uz/v2/class?class={class_id}&view=members&group=parents")
            class_soup = BeautifulSoup(r.content, "html.parser")
            all_users_in_class = class_soup.find_all("tr")
            for class_user in all_users_in_class:
                try:
                    anker = class_user.find(class_="user_ava").find("a")
                    class_user_id = anker.get("href").split("user=")[-1]
                    sex = anker.get("class")[-1]
                    farzand = class_user.find_all("td")[2].find("a").get_text().strip()
                    if class_user_id in natija:
                        natija[class_user_id]["sinf"] = all_classes[grade][class_id]
                        natija[class_user_id]["farzand"] = farzand
                        natija[class_user_id]["sex"] = sex
                except:
                    pass
    database.dict_data["all_data_logins"] = natija
    database.dict_data["profile_kundalikcom"] = {"login": login, "parol": parol, "maktab_id": school_id,  "all_classes": get_all_group(browser, school_id), "mal": database.dict_data["profile_kundalikcom"]["mal"]}
    ui.pushButton.setStyleSheet("""QPushButton{
        background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:1, y2:0.409, stop:0 rgba(255, 241, 0), stop:0.517241 rgba(74, 255, 0), stop:1 rgba(0, 232, 255));
        color: black;
        border-radius: 10px;
        padding: 10px;
        border: 2px solid rgb(0,0,0,0);
        }
        QPushButton:hover{
            background-color: rgb(140,140,140,140);
        border-color:  white;
        }""")

    ui.label_3.setText("100%")
    end = value
    ui.label_3.setStyleSheet(f"""
        background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:1, y2:0, stop:0 rgba(219, 255, 0, 255), stop:1.0 rgba(0, 255, 50, 255));
        color: rgb(0,20, 55);
        font: 30px "Showcard Gothic";""")
    ui.pushButton.setEnabled(True)


def login_user(browser, login, password):
    try:
        r = browser.post(
            "https://login.emaktab.uz/",
            {"exceededAttempts": "False", "login": login, "password": password},
        )
        soup = BeautifulSoup(r.content, "html.parser")
        if "Chiqish" in soup.get_text():
            user_id = soup.find(title="Sozlamalar").get("href").split("user=")[-1].strip()
            
            return True, {
                "user_id": user_id,
                "browser": browser
            }
        elif "Выход" in soup.get_text():
            user_id = soup.find(title="Настройки").get("href").split("user=")[-1].strip()
            return True, {
                "user_id": user_id,
                "browser": browser
            }
        else:
            if soup.find_all(class_="message")[0].get_text().strip() == "Parol yoki login notoʻgʻri koʻrsatilgan. Qaytadan urinib koʻring.":
                return False, "Login yoki parol xato"
            else:
                while True:
                    try:
                        capcha_id = unquote(r.cookies.get('sst')).split("|")[0]
                        url = f"https://login.emaktab.uz/captcha/True/{capcha_id}"
                        file = requests.get(url).content
                        kod = render_capcha_image.to_str(file)
                        r = browser.post(
                            "https://login.emaktab.uz/",
                            {"exceededAttempts": "True", "login": login, "password": password, "Captcha.Input": kod, "Captcha.Id": capcha_id},
                        )
                        soup = BeautifulSoup(r.content, "html.parser")
                        if "Chiqish" in soup.get_text():
                            user_id = soup.find(title="Sozlamalar").get("href").split("user=")[-1].strip()
                            return True, {
                                "user_id": user_id,
                                "browser": browser
                            }
                        elif "Выход" in soup.get_text():
                            user_id = soup.find(title="Настройки").get("href").split("user=")[-1].strip()
                            return True, {
                                "user_id": user_id,
                                "browser": browser
                            }
                        elif soup.find_all(class_="message")[0].get_text().strip() == "Parol yoki login notoʻgʻri koʻrsatilgan. Qaytadan urinib koʻring." or soup.find_all(class_="message")[0].get_text().strip() == "Неправильно указан пароль или логин. Попробуйте еще раз.":
                            return False, "Login yoki parol xato"
                    except requests.exceptions.ConnectionError:
                        return False, "Internega ulanib bo'lmadi"
                    except Exception as e:
                        print(type(e).__name__)
                        traceback.print_exc()
                        return False, "Profilaktika"
    except requests.exceptions.ConnectionError:
        return False, "Internega ulanib bo'lmadi"
    except:
        return False, "Xatolik"



def set_default_font(paragraph, font_name='Times New Roman', font_size=16):
    
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
def no_parol_users(all_data, loading_quit):
    paroli_xatolar = dict()
    for user_id in all_data.keys():
        if all_data[user_id]["login"] == None or all_data[user_id]["parol"] == None:
            if all_data[user_id]["sinf"] not in paroli_xatolar:
                paroli_xatolar[all_data[user_id]["sinf"]] = dict()
            paroli_xatolar[all_data[user_id]["sinf"]][user_id] = all_data[user_id]
    doc = Document()
    sinflar = list(paroli_xatolar.keys())
    sinflar.sort()
    for sinf in sinflar:
        if sinf == "Aniqlanmagan":
            h = doc.add_heading(f"Sinfi aniqlanmagan o'quvchilar hamda ota onalar ichida paroli kiritilmaganlari ro'yxati", level=1)
        else:
            h = doc.add_heading(f"{sinf} sinfdan paroli aniqlanmaganlar ro'yxati", level=1)
        h.style.font.underline = True
        h.style.font.size = Pt(24)
        p = doc.add_paragraph(f"Jami: {len(paroli_xatolar[sinf])} ta")
        p_format = p.paragraph_format
        p_format.space_before = 0
        p_format.space_after = 0
        set_default_font(p)
        n=0
        for user_id in paroli_xatolar[sinf].keys():
            n+=1
            if "farzand" not in all_data[user_id]:
                p = doc.add_paragraph(f"{n}) "+all_data[user_id]["full_name"])
            elif all_data[user_id]["sex"] == "male":
                p = doc.add_paragraph(f"{n}) "+all_data[user_id]["full_name"]+f"\n    >>> {all_data[user_id]['farzand']} ning otasi")
            else:
                p = doc.add_paragraph(f"{n}) "+all_data[user_id]["full_name"]+f"\n    >>> {all_data[user_id]['farzand']} ning onasi")
            p_format = p.paragraph_format
            p_format.space_before = 0
            p_format.space_after = 0
            set_default_font(p)
    # Temporar (vaqtincha) fayl yaratish
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        temp_file_name = tmp_file.name
        doc.save(temp_file_name)
    
    # Hujjatni ochish
    os.startfile(temp_file_name)
    loading_quit()
    # os.remove(temp_file_name)
# no_parol_users({
#     '1000016764321': {
#         'id': '1000016764321',
#         'full_name': 'Chulliyev Jumaboy Rahmanberdiyivich',
#         'sinf': 'Aniqlanmagan',
#         'login': None,
#         'password': None,
#         'parol': None},
#     '1000009405645': {
#         'id': '1000009405645',
#         'full_name': 'Chulliyev Maqsud Ravshanovich',
#         'sinf': '10-a',
#         'login': None,
#         'password': None,
#         'parol': None
#     }
# })

# _, data = login_user(requests.session(), "shamsiddinova.f02198", "1985Feruza")
# print(data)